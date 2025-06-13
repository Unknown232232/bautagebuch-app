# routes/wochenbericht.py
from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, make_response
from flask_login import login_required, current_user
from app import db
from app.models.bautagebuch import BautagebuchEntry, WochenExport
from app.models.aufmass import AufmassEntry
from datetime import datetime, date, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import io
import os

wochenbericht_bp = Blueprint('wochenbericht', __name__)

@wochenbericht_bp.route('/wochenbericht')
@login_required
def wochenbericht_uebersicht():
    """Übersicht aller Wochen mit Exportstatus"""
    
    # Nur Bauleiter und Admin haben Zugriff
    if current_user.role == 'mitarbeiter':
        flash('Keine Berechtigung für Wochenberichte', 'error')
        return redirect(url_for('dashboard.index'))
    
    # Aktuelle Woche
    heute = date.today()
    aktuelle_kw = heute.isocalendar()[1]
    aktuelles_jahr = heute.year
    
    # Alle Wochen mit Einträgen finden
    wochen_query = db.session.query(
        BautagebuchEntry.kalenderwoche,
        BautagebuchEntry.jahr,
        db.func.count(BautagebuchEntry.id).label('anzahl_eintraege')
    ).group_by(
        BautagebuchEntry.kalenderwoche,
        BautagebuchEntry.jahr
    ).order_by(
        BautagebuchEntry.jahr.desc(),
        BautagebuchEntry.kalenderwoche.desc()
    ).all()
    
    # Export-Status ermitteln
    wochen_mit_status = []
    for kw, jahr, anzahl in wochen_query:
        export = WochenExport.query.filter_by(
            kalenderwoche=kw,
            jahr=jahr
        ).first()
        
        # Wochendaten berechnen
        montag = datetime.strptime(f'{jahr}-W{kw:02d}-1', "%Y-W%W-%w").date()
        freitag = montag + timedelta(days=4)
        
        wochen_mit_status.append({
            'kalenderwoche': kw,
            'jahr': jahr,
            'anzahl_eintraege': anzahl,
            'exportiert': export is not None,
            'export_datum': export.exportiert_am if export else None,
            'exportiert_von': export.exportiert_von if export else None,
            'montag': montag,
            'freitag': freitag
        })
    
    return render_template('wochenbericht/uebersicht.html', 
                         wochen=wochen_mit_status,
                         aktuelle_kw=aktuelle_kw,
                         aktuelles_jahr=aktuelles_jahr)

@wochenbericht_bp.route('/wochenbericht/<int:jahr>/<int:kw>')
@login_required
def wochenbericht_anzeigen(jahr, kw):
    """Einzelnen Wochenbericht anzeigen"""
    
    if current_user.role == 'mitarbeiter':
        flash('Keine Berechtigung für Wochenberichte', 'error')
        return redirect(url_for('dashboard.index'))
    
    # Einträge der Woche laden
    eintraege = BautagebuchEntry.query.filter_by(
        kalenderwoche=kw,
        jahr=jahr
    ).order_by(BautagebuchEntry.datum).all()
    
    if not eintraege:
        flash(f'Keine Einträge für KW {kw}/{jahr} gefunden', 'warning')
        return redirect(url_for('wochenbericht.wochenbericht_uebersicht'))
    
    # Wochendaten
    montag = datetime.strptime(f'{jahr}-W{kw:02d}-1', "%Y-W%W-%w").date()
    freitag = montag + timedelta(days=4)
    
    # Export-Status
    export = WochenExport.query.filter_by(kalenderwoche=kw, jahr=jahr).first()
    
    # Statistiken
    stats = {
        'gesamt_eintraege': len(eintraege),
        'mitarbeiter': len(set([e.aufmass_ref.mitarbeiter for e in eintraege])),
        'orte': len(set([e.aufmass_ref.ort for e in eintraege])),
        'materialien': len(set([e.aufmass_ref.material for e in eintraege]))
    }
    
    return render_template('wochenbericht/anzeigen.html',
                         eintraege=eintraege,
                         kalenderwoche=kw,
                         jahr=jahr,
                         montag=montag,
                         freitag=freitag,
                         export=export,
                         stats=stats)

@wochenbericht_bp.route('/wochenbericht/<int:jahr>/<int:kw>/export/<format>')
@login_required
def wochenbericht_export(jahr, kw, format):
    """Wochenbericht als PDF oder Word exportieren"""
    
    if current_user.role == 'mitarbeiter':
        flash('Keine Berechtigung für Export', 'error')
        return redirect(url_for('dashboard.index'))
    
    if format not in ['pdf', 'word']:
        flash('Ungültiges Export-Format', 'error')
        return redirect(url_for('wochenbericht.wochenbericht_anzeigen', jahr=jahr, kw=kw))
    
    # Einträge laden
    eintraege = BautagebuchEntry.query.filter_by(
        kalenderwoche=kw,
        jahr=jahr
    ).order_by(BautagebuchEntry.datum).all()
    
    if not eintraege:
        flash('Keine Einträge zum Exportieren', 'error')
        return redirect(url_for('wochenbericht.wochenbericht_uebersicht'))
    
    # Wochendaten
    montag = datetime.strptime(f'{jahr}-W{kw:02d}-1', "%Y-W%W-%w").date()
    freitag = montag + timedelta(days=4)
    
    if format == 'pdf':
        return export_pdf(eintraege, kw, jahr, montag, freitag)
    else:
        return export_word(eintraege, kw, jahr, montag, freitag)

def export_pdf(eintraege, kw, jahr, montag, freitag):
    """PDF Export erstellen"""
    buffer = io.BytesIO()
    
    # PDF erstellen
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Titel
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=1,  # Zentriert
        spaceAfter=20
    )
    
    story.append(Paragraph(f"Bautagebuch - Kalenderwoche {kw}/{jahr}", title_style))
    story.append(Paragraph(f"Zeitraum: {montag.strftime('%d.%m.%Y')} bis {freitag.strftime('%d.%m.%Y')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Einträge nach Datum gruppieren
    eintraege_by_datum = {}
    for eintrag in eintraege:
        datum_str = eintrag.datum.strftime('%d.%m.%Y')
        if datum_str not in eintraege_by_datum:
            eintraege_by_datum[datum_str] = []
        eintraege_by_datum[datum_str].append(eintrag)
    
    # Einträge hinzufügen
    for datum_str, tages_eintraege in eintraege_by_datum.items():
        story.append(Paragraph(f"<b>{datum_str}</b>", styles['Heading2']))
        
        for eintrag in tages_eintraege:
            story.append(Paragraph(f"• {eintrag.text}", styles['Normal']))
        
        story.append(Spacer(1, 10))
    
    # PDF generieren
    doc.build(story)
    buffer.seek(0)
    
    # Export in DB speichern
    speichere_export(kw, jahr, f"bautagebuch_kw{kw}_{jahr}.pdf")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"bautagebuch_kw{kw}_{jahr}.pdf",
        mimetype='application/pdf'
    )

def export_word(eintraege, kw, jahr, montag, freitag):
    """Word Export erstellen"""
    
    # Word Dokument erstellen
    doc = Document()
    
    # Titel
    title = doc.add_heading(f'Bautagebuch - KW {kw}/{jahr}', 0)
    doc.add_paragraph(f'Zeitraum: {montag.strftime("%d.%m.%Y")} bis {freitag.strftime("%d.%m.%Y")}')
    doc.add_paragraph('')
    
    # Einträge nach Datum gruppieren
    eintraege_by_datum = {}
    for eintrag in eintraege:
        datum_str = eintrag.datum.strftime('%d.%m.%Y')
        if datum_str not in eintraege_by_datum:
            eintraege_by_datum[datum_str] = []
        eintraege_by_datum[datum_str].append(eintrag)
    
    # Einträge hinzufügen
    for datum_str, tages_eintraege in eintraege_by_datum.items():
        doc.add_heading(datum_str, level=1)
        
        for eintrag in tages_eintraege:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(eintrag.text)
        
        doc.add_paragraph('')
    
    # In Buffer speichern
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Export in DB speichern
    speichere_export(kw, jahr, f"bautagebuch_kw{kw}_{jahr}.docx")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"bautagebuch_kw{kw}_{jahr}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def speichere_export(kw, jahr, dateiname):
    """Export-Status in Datenbank speichern"""
    
    # Prüfen ob bereits exportiert
    existing = WochenExport.query.filter_by(
        kalenderwoche=kw,
        jahr=jahr
    ).first()
    
    if not existing:
        export = WochenExport(
            kalenderwoche=kw,
            jahr=jahr,
            exportiert_von=current_user.username,
            dateiname=dateiname
        )
        db.session.add(export)
        db.session.commit()
        
        flash(f'Wochenbericht KW {kw}/{jahr} erfolgreich exportiert!', 'success')
